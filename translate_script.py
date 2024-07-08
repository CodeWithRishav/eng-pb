from googletrans import Translator
from docx import Document
import os
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def translate_text(text, target_language="pa"):
    translator = Translator()
    try:
        logging.info(f"Translating text: {text[:30]}...")  # Log start of translation
        translated = translator.translate(text, dest=target_language)
        translated_text = translated.text
        logging.info("Translation successful")
        return translated_text
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        return text  # Return original text if translation fails

def copy_formatting(source_para, target_para):
    target_para.style = source_para.style
    
    # Copy paragraph formatting
    target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
    target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing

    # Copy run formatting
    for src_run, tgt_run in zip(source_para.runs, target_para.runs):
        tgt_run.bold = src_run.bold
        tgt_run.italic = src_run.italic
        tgt_run.underline = src_run.underline
        tgt_run.font.size = src_run.font.size
        tgt_run.font.name = src_run.font.name
        tgt_run.font.color.rgb = src_run.font.color.rgb

def translate_document(input_file, output_file, target_language="pa"):
    doc = Document(input_file)
    translated_doc = Document()

    for para in doc.paragraphs:
        translated_text = translate_text(para.text, target_language)
        translated_para = translated_doc.add_paragraph(translated_text)
        copy_formatting(para, translated_para)

    translated_doc.save(output_file)
    logging.info(f"Document saved to {output_file}")

def translate_documents_in_folder(input_folder, output_folder, target_language="pa"):
    # Ensure output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Iterate over files in input folder
    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):  # Assuming all files are .docx format
            input_file = os.path.join(input_folder, filename)
            output_file = os.path.join(output_folder, filename)
            
            # Translate and save each document
            translate_document(input_file, output_file, target_language)
            logging.info(f"Translated and saved: {output_file}")

# Example usage:
input_folder = r"C:\Users\HP\Downloads\translator\input"
output_folder = r"C:\Users\HP\Downloads\translator\output"
translate_documents_in_folder(input_folder, output_folder)
